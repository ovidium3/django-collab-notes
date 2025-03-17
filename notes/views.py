import os
import re
import mimetypes
from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.conf import settings
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from .models import Note, Highlight, Comment
import json
import PyPDF2
import docx
import uuid

def note_list(request):
    notes = Note.objects.all().order_by('-updated_at')
    return render(request, 'notes/note_list.html', {'notes': notes})

def signup_view(request):
    if request.method == "POST":
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('/')  # Redirect to homepage
    else:
        form = UserCreationForm()
    return render(request, 'signup.html', {'form': form})

@login_required
def create_note(request):
    if request.method == 'POST':
        if request.FILES.get('file'):
            # handle file upload
            uploaded_file = request.FILES['file']
            file_content = ""
            original_filename = uploaded_file.name
            
            # extract content based on file type
            if uploaded_file.name.endswith('.pdf'):
                # process PDF
                temp_file_path = default_storage.save('temp/' + uploaded_file.name, ContentFile(uploaded_file.read()))
                try:
                    with open(default_storage.path(temp_file_path), 'rb') as file:
                        pdf_reader = PyPDF2.PdfFileReader(file)
                        for page_num in range(pdf_reader.numPages):
                            page = pdf_reader.getPage(page_num)
                            file_content += page.extractText() + "\n\n"
                finally:
                    default_storage.delete(temp_file_path)
            
            elif uploaded_file.name.endswith('.docx'):
                # process Word doc
                temp_file_path = default_storage.save('temp/' + uploaded_file.name, ContentFile(uploaded_file.read()))
                try:
                    doc = docx.Document(default_storage.path(temp_file_path))
                    file_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                finally:
                    default_storage.delete(temp_file_path)
            
            elif uploaded_file.name.endswith('.txt'):
                # process text file
                file_content = uploaded_file.read().decode('utf-8')
            
            else:
                return JsonResponse({'error': 'Unsupported file format. Please upload PDF, DOCX, or TXT files.'})
            
            # clean up content
            file_content = re.sub(r'\s+', ' ', file_content).strip()
            
            # create note with the extracted content
            title = request.POST.get('title', os.path.splitext(original_filename)[0])
            note = Note.objects.create(
                title=title,
                content=file_content,
                is_uploaded=True,
                original_filename=original_filename
            )
            
            if 'application/json' in request.headers.get('Accept', ''):
                return JsonResponse({'success': True, 'note_id': note.id})
            else:
                return redirect('editor', note_id=note.id)
        
        elif request.headers.get('Content-Type') == 'application/json':
            # handle JSON data
            data = json.loads(request.body)
            note = Note.objects.create(
                title=data.get('title', 'Untitled Note'),
                content=data.get('content', 'This is a collaborative note. Start editing!')
            )
            return JsonResponse({'success': True, 'note_id': note.id})
        
        else:
            # handle form data
            title = request.POST.get('title', 'Untitled Note')
            content = request.POST.get('content', 'This is a collaborative note. Start editing!')
            note = Note.objects.create(title=title, content=content)
            return redirect('editor', note_id=note.id)
    
    return render(request, 'notes/create_note.html')

@csrf_exempt
def delete_note(request, note_id):
    note = get_object_or_404(Note, id=note_id)
    if request.method == 'POST':
        note.delete()
        return JsonResponse({'success': True})
    return JsonResponse({'success': False}, status=400)

def editor(request, note_id):
    note = get_object_or_404(Note, id=note_id)
    highlights = note.highlights.all()
    
    # gather comments for each highlight
    highlight_comments = {}
    for highlight in highlights:
        highlight_comments[highlight.id] = list(highlight.comments.values('id', 'text'))
    
    context = {
        'note': note,
        'highlights': list(highlights.values('id', 'start_offset', 'end_offset')),
        'highlight_comments': highlight_comments
    }
    
    return render(request, 'notes/editor.html', context)

# File upload handler
@login_required
@require_POST
def upload_file(request):
    if 'file' not in request.FILES:
        return JsonResponse({'success': False, 'error': 'No file was provided'})
    
    file = request.FILES['file']
    note_id = request.POST.get('note_id')
    
    try:
        note = Note.objects.get(id=note_id, user=request.user)
    except Note.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Note not found'})
    
    # Generate a unique filename to prevent conflicts
    filename = f"{uuid.uuid4()}_{file.name}"
    
    # Define the upload path - store in a subdirectory for each note
    upload_path = f"uploads/{note_id}/{filename}"
    
    # Save the file using Django's storage API
    file_path = default_storage.save(upload_path, ContentFile(file.read()))
    
    # Generate the URL for the file
    file_url = default_storage.url(file_path)
    
    return JsonResponse({
        'success': True, 
        'file_url': file_url,
        'filename': file.name
    })

# Document import handler
@login_required
@require_POST
def import_document(request):
    if 'file' not in request.FILES:
        return JsonResponse({'success': False, 'error': 'No file was provided'})
    
    file = request.FILES['file']
    title = request.POST.get('title', file.name)
    
    # Check file size (limit to 10MB)
    if file.size > 10 * 1024 * 1024:  # 10MB in bytes
        return JsonResponse({'success': False, 'error': 'File is too large. Maximum size is 10MB'})
    
    # Process different file types
    file_extension = os.path.splitext(file.name)[1].lower()
    
    content = ""
    
    try:
        if file_extension == '.txt':
            # Process text file
            content = file.read().decode('utf-8')
        
        elif file_extension == '.md':
            # Process markdown file
            content = file.read().decode('utf-8')
        
        elif file_extension == '.docx':
            # Process Word document (requires python-docx)
            try:
                import docx
                doc = docx.Document(file)
                content = '\n\n'.join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                return JsonResponse({'success': False, 'error': 'DOCX processing is not available'})
        
        elif file_extension == '.pdf':
            # Process PDF (requires PyPDF2)
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file)
                content = '\n\n'.join([page.extract_text() for page in reader.pages])
            except ImportError:
                return JsonResponse({'success': False, 'error': 'PDF processing is not available'})
        
        else:
            return JsonResponse({'success': False, 'error': 'Unsupported file format'})
        
        # Create new note with the imported content
        note = Note.objects.create(
            user=request.user,
            title=title,
            content=content
        )
        
        # Store the original file for reference
        filename = f"{uuid.uuid4()}_{file.name}"
        upload_path = f"uploads/{note.id}/{filename}"
        file_path = default_storage.save(upload_path, ContentFile(file.read()))
        
        # Add a reference to the original file in the note content
        file_url = default_storage.url(file_path)
        file_ref = f"\n\n[Original file: {file.name}]({file_url})\n"
        note.content += file_ref
        note.save()
        
        return JsonResponse({
            'success': True,
            'note_id': note.id
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

# Update note content
@login_required
@require_POST
def update_note_content(request, note_id):
    try:
        note = Note.objects.get(id=note_id, user=request.user)
    except Note.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Note not found'})
    
    try:
        data = json.loads(request.body)
        content = data.get('content')
        
        if content is not None:
            note.content = content
            note.save()
            
            return JsonResponse({'success': True})
        else:
            return JsonResponse({'success': False, 'error': 'No content provided'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})
